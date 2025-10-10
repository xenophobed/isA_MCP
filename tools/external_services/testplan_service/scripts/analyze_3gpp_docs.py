#!/usr/bin/env python3
"""
Analyze 3GPP documents to understand structure and identify extraction points
"""

import os
from pathlib import Path
import re
from typing import Dict, List, Any
import json
import sys

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

try:
    from python_docx import Document
except ImportError:
    try:
        from docx import Document
    except ImportError:
        print("Warning: python-docx not available, using fallback")
        Document = None

class ThreeGPPDocumentAnalyzer:
    def __init__(self):
        self.base_dir = Path("/Users/xenodennis/Documents/Fun/isA_MCP/tools/external_services/testplan_service")
        self.data_dir = self.base_dir / "data_source" / "3GPP"
        
        # Key specifications to analyze (mapping from PICS to test cases)
        self.spec_mappings = {
            'TS 36.521-2': 'TS 36.521-1',  # PICS → RF conformance tests
            'TS 36.523-2': 'TS 36.523-1',  # PICS → Protocol conformance tests
            'TS 38.508-2': 'TS 38.523-1',  # PICS → 5G protocol tests
            'TS 34.123-2': 'TS 34.123-1',  # PICS → 3G protocol tests
        }
        
        self.analysis_results = {}
    
    def extract_text_from_docx(self, doc_path: Path) -> str:
        """Extract text from DOCX file using XML parsing"""
        import zipfile
        from xml.etree import ElementTree
        
        text_content = []
        try:
            with zipfile.ZipFile(str(doc_path), 'r') as docx:
                # Read document.xml
                xml_content = docx.read('word/document.xml')
                tree = ElementTree.fromstring(xml_content)
                
                # Extract all text
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for paragraph in tree.findall('.//w:p', namespaces):
                    texts = []
                    for node in paragraph.findall('.//w:t', namespaces):
                        if node.text:
                            texts.append(node.text)
                    if texts:
                        text_content.append(' '.join(texts))
        except Exception as e:
            print(f"Error extracting from {doc_path}: {e}")
        
        return '\n'.join(text_content)
    
    def _analyze_text_content(self, content: str, doc_path: Path) -> Dict[str, Any]:
        """Analyze text content directly"""
        lines = content.split('\n')
        
        analysis = {
            'file': doc_path.name,
            'spec': doc_path.parent.parent.name,
            'paragraphs': len(lines),
            'tables': 0,  # Can't detect tables from plain text easily
            'sections': [],
            'test_cases': [],
            'pics_references': [],
            'sample_content': []
        }
        
        for i, line in enumerate(lines):
            text = line.strip()
            if not text:
                continue
            
            # Detect section headers (usually numbered)
            if re.match(r'^(\d+\.?)+\s+', text):
                if i < 100:  # Only early sections
                    analysis['sections'].append(text[:100])
            
            # Detect test case definitions
            if any(keyword in text.lower() for keyword in ['test case', 'test purpose', 'test procedure']):
                analysis['test_cases'].append(text[:200])
            
            # Detect PICS references
            if 'PICS' in text or re.search(r'[A-Z]\.\d+[\./]\d+', text):
                analysis['pics_references'].append(text[:150])
            
            # Sample early content
            if i < 10 and text:
                analysis['sample_content'].append(text[:200])
        
        return analysis
    
    def _analyze_with_docx(self, doc, doc_path: Path) -> Dict[str, Any]:
        """Analyze using python-docx Document object"""
        analysis = {
            'file': doc_path.name,
            'spec': doc_path.parent.parent.name,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': [],
            'test_cases': [],
            'pics_references': [],
            'sample_content': []
        }
        
        # Analyze paragraphs for structure
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Detect section headers (usually numbered)
            if re.match(r'^(\d+\.?)+\s+', text):
                if i < 50:  # Only early sections
                    analysis['sections'].append(text[:100])
            
            # Detect test case definitions
            if any(keyword in text.lower() for keyword in ['test case', 'test purpose', 'test procedure']):
                analysis['test_cases'].append(text[:200])
            
            # Detect PICS references
            if 'PICS' in text or re.search(r'[A-Z]\.\d+[\./]\d+', text):
                analysis['pics_references'].append(text[:150])
            
            # Sample early content
            if i < 10 and text:
                analysis['sample_content'].append(text[:200])
        
        # Analyze tables (often contain test case details)
        for table_idx, table in enumerate(doc.tables):
            if table_idx >= 5:  # Analyze first 5 tables
                break
            
            # Get headers from first row
            if len(table.rows) > 0:
                headers = []
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip()[:50])
                
                # Look for test-related tables
                if any('test' in h.lower() for h in headers):
                    analysis['test_cases'].append(f"Table {table_idx}: {', '.join(headers)}")
                
                # Look for PICS-related tables
                if any('pics' in h.lower() or 'applicability' in h.lower() for h in headers):
                    analysis['pics_references'].append(f"Table {table_idx}: {', '.join(headers)}")
        
        return analysis
    
    def analyze_document_structure(self, doc_path: Path) -> Dict[str, Any]:
        """Analyze a single document's structure"""
        try:
            # Try using Document if available
            if Document:
                doc = Document(str(doc_path))
                return self._analyze_with_docx(doc, doc_path)
            else:
                # Fallback to XML parsing
                content = self.extract_text_from_docx(doc_path)
                return self._analyze_text_content(content, doc_path)
            
        except Exception as e:
            return {'error': str(e), 'file': str(doc_path)}
    
    def analyze_test_specification(self, spec_name: str):
        """Analyze a test specification (e.g., 36.523-1)"""
        print(f"\n{'='*70}")
        print(f"Analyzing {spec_name}")
        print(f"{'='*70}")
        
        spec_dir = self.data_dir / spec_name
        if not spec_dir.exists():
            print(f"Specification directory not found: {spec_dir}")
            return
        
        # Find all docx files
        doc_files = list(spec_dir.rglob("*.docx"))
        print(f"Found {len(doc_files)} documents")
        
        spec_analysis = {
            'specification': spec_name,
            'total_files': len(doc_files),
            'documents': []
        }
        
        # Analyze each document
        for doc_file in sorted(doc_files)[:3]:  # Analyze first 3 docs
            print(f"\nAnalyzing: {doc_file.name}")
            doc_analysis = self.analyze_document_structure(doc_file)
            
            # Print summary
            print(f"  - Paragraphs: {doc_analysis.get('paragraphs', 0)}")
            print(f"  - Tables: {doc_analysis.get('tables', 0)}")
            print(f"  - Sections found: {len(doc_analysis.get('sections', []))}")
            print(f"  - Test cases found: {len(doc_analysis.get('test_cases', []))}")
            print(f"  - PICS references: {len(doc_analysis.get('pics_references', []))}")
            
            # Show samples
            if doc_analysis.get('sections'):
                print(f"\n  Sample sections:")
                for section in doc_analysis['sections'][:3]:
                    print(f"    • {section}")
            
            if doc_analysis.get('test_cases'):
                print(f"\n  Sample test cases:")
                for tc in doc_analysis['test_cases'][:3]:
                    print(f"    • {tc}")
            
            if doc_analysis.get('pics_references'):
                print(f"\n  Sample PICS references:")
                for pics in doc_analysis['pics_references'][:3]:
                    print(f"    • {pics}")
            
            spec_analysis['documents'].append(doc_analysis)
        
        self.analysis_results[spec_name] = spec_analysis
        return spec_analysis
    
    def identify_extraction_patterns(self):
        """Identify common patterns for data extraction"""
        print("\n" + "="*70)
        print("EXTRACTION PATTERNS IDENTIFIED")
        print("="*70)
        
        patterns = {
            'test_case_patterns': [
                r'(\d+\.)+\d+\s+Test Case.*',  # Section headers for test cases
                r'Test\s+Case\s+ID:\s*(.*)',    # Test case ID
                r'Test\s+Purpose:\s*(.*)',      # Test purpose
                r'PICS\s+Selection:\s*(.*)',    # PICS requirements
                r'Applicability:\s*(.*)',       # Applicability conditions
            ],
            'pics_patterns': [
                r'[A-Z]\.\d+[\./]\d+[-/]\d+',   # PICS item format (e.g., A.4.3-1)
                r'Table\s+[A-Z]\.\d+.*',        # PICS table references
                r'PICS:\s*(.*)',                # Direct PICS references
            ],
            'table_headers': [
                'Test Case ID',
                'Test Purpose', 
                'PICS Selection',
                'Applicability',
                'Test Case Name',
                'Clause',
                'Release'
            ]
        }
        
        print("\nTest Case Extraction Patterns:")
        for pattern in patterns['test_case_patterns']:
            print(f"  • {pattern}")
        
        print("\nPICS Extraction Patterns:")
        for pattern in patterns['pics_patterns']:
            print(f"  • {pattern}")
        
        print("\nKey Table Headers to Look For:")
        for header in patterns['table_headers']:
            print(f"  • {header}")
        
        return patterns
    
    def analyze_pics_to_test_mapping(self):
        """Analyze how PICS specifications map to test specifications"""
        print("\n" + "="*70)
        print("PICS TO TEST SPECIFICATION MAPPING")
        print("="*70)
        
        for pics_spec, test_spec in self.spec_mappings.items():
            print(f"\n{pics_spec} (PICS) → {test_spec} (Tests)")
            
            # Check if both exist
            pics_dir = self.data_dir / pics_spec
            test_dir = self.data_dir / test_spec
            
            if pics_dir.exists():
                pics_files = list(pics_dir.rglob("*.docx"))
                print(f"  PICS spec: {len(pics_files)} documents found")
            else:
                print(f"  PICS spec: NOT FOUND")
            
            if test_dir.exists():
                test_files = list(test_dir.rglob("*.docx"))
                print(f"  Test spec: {len(test_files)} documents found")
            else:
                print(f"  Test spec: NOT FOUND")
    
    def run_analysis(self):
        """Run complete analysis"""
        # Analyze PICS to test mappings
        self.analyze_pics_to_test_mapping()
        
        # Analyze key test specifications
        test_specs = ['TS 36.523-1', 'TS 36.521-1', 'TS 38.523-1']
        
        for spec in test_specs:
            if (self.data_dir / spec).exists():
                self.analyze_test_specification(spec)
        
        # Identify extraction patterns
        patterns = self.identify_extraction_patterns()
        
        # Save analysis results
        output_file = self.base_dir / "3gpp_document_analysis.json"
        with open(output_file, 'w') as f:
            json.dump({
                'analysis_results': self.analysis_results,
                'extraction_patterns': patterns,
                'spec_mappings': self.spec_mappings
            }, f, indent=2, default=str)
        
        print(f"\n✓ Analysis saved to: {output_file}")
        
        # Summary
        print("\n" + "="*70)
        print("ANALYSIS SUMMARY")
        print("="*70)
        print("\n📊 Key Findings:")
        print("1. Test specifications contain test cases with PICS requirements")
        print("2. Test cases are organized in sections and tables")
        print("3. Each test case has ID, purpose, and PICS selection criteria")
        print("4. PICS items follow pattern: A.4.3-1, B.2.1/5, etc.")
        print("5. Tables often contain structured test case information")
        
        print("\n🎯 Data Extraction Requirements:")
        print("1. Extract test case ID, name, and purpose")
        print("2. Extract PICS requirements for each test")
        print("3. Extract applicability conditions")
        print("4. Build mapping: PICS item → applicable test cases")
        print("5. Store in database for efficient querying")


if __name__ == "__main__":
    analyzer = ThreeGPPDocumentAnalyzer()
    analyzer.run_analysis()