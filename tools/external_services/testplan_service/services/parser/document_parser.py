#!/usr/bin/env python3
"""
Universal Document Parser for 3GPP Specifications

Handles multiple document formats and provides unified access to:
- Tables
- Paragraphs
- Document structure
- Multi-file merging

Supported formats:
- .docx (Word 2007+)
- .doc (Word 97-2003)
- .pdf (PDF documents)
- .xlsx, .xlsm (Excel files)
- .txt (Plain text)
"""

import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from datetime import datetime

logger = logging.getLogger(__name__)


@dataclass
class Table:
    """Represents a table from a document"""
    index: int
    headers: List[str]
    rows: List[List[str]]
    page: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert table to dictionary format"""
        return {
            'index': self.index,
            'headers': self.headers,
            'rows': self.rows,
            'page': self.page,
            'metadata': self.metadata
        }
    
    def get_column(self, header: str) -> List[str]:
        """Get all values from a specific column"""
        if header not in self.headers:
            return []
        idx = self.headers.index(header)
        return [row[idx] if idx < len(row) else '' for row in self.rows]


@dataclass
class ParsedDocument:
    """Result of document parsing"""
    success: bool
    file_path: Path
    file_type: str
    tables: List[Table]
    paragraphs: List[str]
    metadata: Dict[str, Any]
    error: Optional[str] = None
    
    def merge_with(self, other: 'ParsedDocument'):
        """Merge another document's content into this one"""
        self.tables.extend(other.tables)
        self.paragraphs.extend(other.paragraphs)
        self.metadata['merged_files'] = self.metadata.get('merged_files', []) + [str(other.file_path)]


class UniversalDocumentParser:
    """
    Universal parser supporting multiple document formats
    """
    
    SUPPORTED_FORMATS = {
        '.docx', '.doc', '.pdf', '.xlsx', '.xlsm', '.xls', '.txt'
    }
    
    def __init__(self):
        """Initialize the parser with available libraries"""
        self.available_libraries = self._check_available_libraries()
        logger.info(f"Available libraries: {self.available_libraries}")
    
    def _check_available_libraries(self) -> Dict[str, bool]:
        """Check which document processing libraries are available"""
        libraries = {}
        
        # Check for python-docx
        try:
            import docx
            libraries['python-docx'] = True
        except ImportError:
            libraries['python-docx'] = False
        
        # Check for docx2txt
        try:
            import docx2txt
            libraries['docx2txt'] = True
        except ImportError:
            libraries['docx2txt'] = False
        
        # Check for mammoth
        try:
            import mammoth
            libraries['mammoth'] = True
        except ImportError:
            libraries['mammoth'] = False
        
        # Check for doc2docx
        try:
            from doc2docx import convert
            libraries['doc2docx'] = True
        except ImportError:
            libraries['doc2docx'] = False
        
        # Check for pdfplumber
        try:
            import pdfplumber
            libraries['pdfplumber'] = True
        except ImportError:
            libraries['pdfplumber'] = False
        
        # Check for PyPDF2/pypdf
        try:
            import pypdf
            libraries['pypdf'] = True
        except ImportError:
            try:
                import PyPDF2
                libraries['PyPDF2'] = True
            except ImportError:
                libraries['pypdf'] = False
                libraries['PyPDF2'] = False
        
        # Check for pdf2docx
        try:
            import pdf2docx
            libraries['pdf2docx'] = True
        except ImportError:
            libraries['pdf2docx'] = False
        
        # Check for openpyxl
        try:
            import openpyxl
            libraries['openpyxl'] = True
        except ImportError:
            libraries['openpyxl'] = False
        
        # Check for LibreOffice
        try:
            result = subprocess.run(['soffice', '--version'], capture_output=True, text=True)
            libraries['libreoffice'] = result.returncode == 0
        except FileNotFoundError:
            libraries['libreoffice'] = False
        
        return libraries
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse a document file
        
        Args:
            file_path: Path to the document
            
        Returns:
            ParsedDocument with extracted content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type='unknown',
                tables=[],
                paragraphs=[],
                metadata={},
                error=f"File not found: {file_path}"
            )
        
        file_type = file_path.suffix.lower()
        
        if file_type not in self.SUPPORTED_FORMATS:
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type=file_type,
                tables=[],
                paragraphs=[],
                metadata={},
                error=f"Unsupported file format: {file_type}"
            )
        
        logger.info(f"Parsing {file_type} file: {file_path.name}")
        
        # Route to appropriate parser
        if file_type == '.docx':
            return self._parse_docx(file_path)
        elif file_type == '.doc':
            return self._parse_doc(file_path)
        elif file_type == '.pdf':
            return self._parse_pdf(file_path)
        elif file_type in ['.xlsx', '.xlsm', '.xls']:
            return self._parse_excel(file_path)
        elif file_type == '.txt':
            return self._parse_text(file_path)
        else:
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type=file_type,
                tables=[],
                paragraphs=[],
                metadata={},
                error=f"No parser for format: {file_type}"
            )
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse a .docx file"""
        if not self.available_libraries.get('python-docx'):
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type='.docx',
                tables=[],
                paragraphs=[],
                metadata={},
                error="python-docx not installed"
            )
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(row_data)
                
                if table_data:
                    headers = table_data[0] if table_data else []
                    rows = table_data[1:] if len(table_data) > 1 else []
                    
                    tables.append(Table(
                        index=idx,
                        headers=headers,
                        rows=rows,
                        metadata={'source': 'docx'}
                    ))
            
            return ParsedDocument(
                success=True,
                file_path=file_path,
                file_type='.docx',
                tables=tables,
                paragraphs=paragraphs,
                metadata={
                    'parser': 'python-docx',
                    'table_count': len(tables),
                    'paragraph_count': len(paragraphs)
                }
            )
            
        except Exception as e:
            logger.error(f"Failed to parse .docx: {e}")
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type='.docx',
                tables=[],
                paragraphs=[],
                metadata={},
                error=str(e)
            )
    
    def _parse_doc(self, file_path: Path) -> ParsedDocument:
        """Parse a .doc file using multiple fallback methods"""
        
        # Method 1: Try doc2docx conversion (best for cross-platform)
        if self.available_libraries.get('doc2docx'):
            try:
                from doc2docx import convert
                # Create temp output path
                temp_docx = file_path.parent / f"{file_path.stem}_temp.docx"
                
                # Convert .doc to .docx
                convert(str(file_path), str(temp_docx))
                
                if temp_docx.exists():
                    # Parse the converted docx file
                    result = self._parse_docx(temp_docx)
                    # Clean up temp file
                    temp_docx.unlink()
                    
                    # Update metadata
                    result.metadata['original_format'] = '.doc'
                    result.metadata['conversion_method'] = 'doc2docx'
                    return result
                    
            except Exception as e:
                logger.warning(f"doc2docx conversion failed: {e}")
        
        # Method 2: Try LibreOffice conversion
        if self.available_libraries.get('libreoffice'):
            try:
                converted = self._convert_with_libreoffice(file_path, 'docx')
                if converted:
                    result = self._parse_docx(converted)
                    # Clean up temp file
                    if converted.exists():
                        converted.unlink()
                    return result
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}")
        
        # Method 3: Try mammoth for .doc conversion
        if self.available_libraries.get('mammoth'):
            try:
                import mammoth
                with open(file_path, 'rb') as doc_file:
                    # Extract as plain text
                    result = mammoth.extract_raw_text(doc_file)
                    text = result.value
                    
                # Parse text into paragraphs and tables
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                tables = self._extract_tables_from_text(text)
                
                return ParsedDocument(
                    success=True,
                    file_path=file_path,
                    file_type='.doc',
                    tables=tables,
                    paragraphs=paragraphs,
                    metadata={
                        'parser': 'mammoth',
                        'messages': result.messages if hasattr(result, 'messages') else []
                    }
                )
            except Exception as e:
                logger.warning(f"Mammoth conversion failed: {e}")
        
        # Method 3: Try direct loading with python-docx
        if self.available_libraries.get('python-docx'):
            try:
                from docx import Document
                doc = Document(str(file_path))
                # If successful, process as docx
                return self._parse_docx(file_path)
            except Exception as e:
                logger.warning(f"Direct python-docx loading failed: {e}")
        
        # Method 4: Use docx2txt for text extraction
        if self.available_libraries.get('docx2txt'):
            try:
                import docx2txt
                text = docx2txt.process(str(file_path))
                
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                tables = self._extract_tables_from_text(text)
                
                return ParsedDocument(
                    success=True,
                    file_path=file_path,
                    file_type='.doc',
                    tables=tables,
                    paragraphs=paragraphs,
                    metadata={
                        'parser': 'docx2txt',
                        'warning': 'Tables extracted from text may be incomplete'
                    }
                )
            except Exception as e:
                logger.warning(f"docx2txt extraction failed: {e}")
        
        # Method 4: Last resort - try as text
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                
                return ParsedDocument(
                    success=True,
                    file_path=file_path,
                    file_type='.doc',
                    tables=[],
                    paragraphs=paragraphs,
                    metadata={
                        'parser': 'text_fallback',
                        'warning': 'Binary file read as text, content may be corrupted'
                    }
                )
        except Exception as e:
            logger.error(f"All .doc parsing methods failed: {e}")
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type='.doc',
                tables=[],
                paragraphs=[],
                metadata={},
                error="All parsing methods failed"
            )
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse a PDF file using multiple methods"""
        
        # Method 1: Use pdfplumber (best for tables)
        if self.available_libraries.get('pdfplumber'):
            try:
                import pdfplumber
                
                paragraphs = []
                tables = []
                
                with pdfplumber.open(str(file_path)) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        text = page.extract_text()
                        if text:
                            page_paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                            paragraphs.extend(page_paragraphs)
                        
                        # Extract tables
                        page_tables = page.extract_tables()
                        for table_data in page_tables:
                            if table_data and len(table_data) > 1:
                                # Clean table data
                                cleaned_table = []
                                for row in table_data:
                                    if row:  # Skip None rows
                                        cleaned_row = [str(cell) if cell else '' for cell in row]
                                        if any(cleaned_row):  # Skip empty rows
                                            cleaned_table.append(cleaned_row)
                                
                                if cleaned_table:
                                    headers = cleaned_table[0]
                                    rows = cleaned_table[1:]
                                    
                                    tables.append(Table(
                                        index=len(tables),
                                        headers=headers,
                                        rows=rows,
                                        page=page_num + 1,
                                        metadata={'source': 'pdfplumber'}
                                    ))
                
                return ParsedDocument(
                    success=True,
                    file_path=file_path,
                    file_type='.pdf',
                    tables=tables,
                    paragraphs=paragraphs,
                    metadata={
                        'parser': 'pdfplumber',
                        'page_count': len(pdf.pages),
                        'table_count': len(tables)
                    }
                )
                
            except Exception as e:
                logger.warning(f"pdfplumber parsing failed: {e}")
        
        # Method 2: Use pdf2docx for conversion
        if self.available_libraries.get('pdf2docx'):
            try:
                from pdf2docx import Converter
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_path = Path(tmp_file.name)
                
                cv = Converter(str(file_path))
                cv.convert(str(tmp_path))
                cv.close()
                
                # Parse converted docx
                result = self._parse_docx(tmp_path)
                
                # Clean up
                if tmp_path.exists():
                    tmp_path.unlink()
                
                result.metadata['original_parser'] = 'pdf2docx'
                return result
                
            except Exception as e:
                logger.warning(f"pdf2docx conversion failed: {e}")
        
        # Method 3: Use PyPDF2/pypdf for text extraction
        pdf_library = None
        if self.available_libraries.get('pypdf'):
            try:
                import pypdf as pdf_library
            except ImportError:
                pass
        elif self.available_libraries.get('PyPDF2'):
            try:
                import PyPDF2 as pdf_library
            except ImportError:
                pass
        
        if pdf_library:
            try:
                paragraphs = []
                
                with open(file_path, 'rb') as file:
                    pdf_reader = pdf_library.PdfReader(file) if hasattr(pdf_library, 'PdfReader') else pdf_library.PdfFileReader(file)
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text:
                            page_paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                            paragraphs.extend(page_paragraphs)
                
                return ParsedDocument(
                    success=True,
                    file_path=file_path,
                    file_type='.pdf',
                    tables=[],
                    paragraphs=paragraphs,
                    metadata={
                        'parser': 'pypdf/PyPDF2',
                        'warning': 'No table extraction available'
                    }
                )
                
            except Exception as e:
                logger.warning(f"PyPDF parsing failed: {e}")
        
        return ParsedDocument(
            success=False,
            file_path=file_path,
            file_type='.pdf',
            tables=[],
            paragraphs=[],
            metadata={},
            error="No PDF parser available"
        )
    
    def _parse_excel(self, file_path: Path) -> ParsedDocument:
        """Parse Excel files"""
        if not self.available_libraries.get('openpyxl'):
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type=file_path.suffix.lower(),
                tables=[],
                paragraphs=[],
                metadata={},
                error="openpyxl not installed"
            )
        
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(str(file_path), data_only=True)
            tables = []
            paragraphs = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # Extract data from sheet
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_data):  # Skip empty rows
                        sheet_data.append(row_data)
                
                if sheet_data:
                    headers = sheet_data[0]
                    rows = sheet_data[1:]
                    
                    tables.append(Table(
                        index=len(tables),
                        headers=headers,
                        rows=rows,
                        metadata={
                            'sheet_name': sheet_name,
                            'source': 'openpyxl'
                        }
                    ))
            
            wb.close()
            
            return ParsedDocument(
                success=True,
                file_path=file_path,
                file_type=file_path.suffix.lower(),
                tables=tables,
                paragraphs=paragraphs,
                metadata={
                    'parser': 'openpyxl',
                    'sheet_count': len(wb.sheetnames),
                    'table_count': len(tables)
                }
            )
            
        except Exception as e:
            logger.error(f"Failed to parse Excel file: {e}")
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type=file_path.suffix.lower(),
                tables=[],
                paragraphs=[],
                metadata={},
                error=str(e)
            )
    
    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            tables = self._extract_tables_from_text(text)
            
            return ParsedDocument(
                success=True,
                file_path=file_path,
                file_type='.txt',
                tables=tables,
                paragraphs=paragraphs,
                metadata={
                    'parser': 'text',
                    'size': len(text)
                }
            )
            
        except Exception as e:
            logger.error(f"Failed to parse text file: {e}")
            return ParsedDocument(
                success=False,
                file_path=file_path,
                file_type='.txt',
                tables=[],
                paragraphs=[],
                metadata={},
                error=str(e)
            )
    
    def _convert_with_libreoffice(self, file_path: Path, target_format: str) -> Optional[Path]:
        """Convert document using LibreOffice"""
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                cmd = [
                    'soffice', '--headless', '--convert-to', target_format,
                    '--outdir', tmpdir, str(file_path)
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    # Find converted file
                    converted_path = Path(tmpdir) / f"{file_path.stem}.{target_format}"
                    if converted_path.exists():
                        # Copy to a persistent location
                        output_path = file_path.parent / f"{file_path.stem}_converted.{target_format}"
                        import shutil
                        shutil.copy2(converted_path, output_path)
                        return output_path
                        
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {e}")
        
        return None
    
    def _extract_tables_from_text(self, text: str) -> List[Table]:
        """Extract tables from plain text using heuristics"""
        tables = []
        lines = text.split('\n')
        
        # Look for table-like patterns
        current_table = []
        
        for line in lines:
            # Check if line looks like a table row (has multiple columns)
            if '\t' in line or re.search(r'\s{2,}', line):
                # Split by tabs or multiple spaces
                cells = re.split(r'\t+|\s{2,}', line.strip())
                if len(cells) > 1:
                    current_table.append(cells)
            elif current_table:
                # End of table
                if len(current_table) > 1:  # At least 2 rows
                    headers = current_table[0]
                    rows = current_table[1:]
                    
                    tables.append(Table(
                        index=len(tables),
                        headers=headers,
                        rows=rows,
                        metadata={'source': 'text_extraction'}
                    ))
                current_table = []
        
        # Add last table if exists
        if len(current_table) > 1:
            headers = current_table[0]
            rows = current_table[1:]
            
            tables.append(Table(
                index=len(tables),
                headers=headers,
                rows=rows,
                metadata={'source': 'text_extraction'}
            ))
        
        return tables
    
    def merge_documents(self, file_paths: List[Path]) -> ParsedDocument:
        """
        Merge multiple documents into a single ParsedDocument
        
        Args:
            file_paths: List of document paths to merge
            
        Returns:
            Merged ParsedDocument
        """
        if not file_paths:
            return ParsedDocument(
                success=False,
                file_path=Path('.'),
                file_type='merged',
                tables=[],
                paragraphs=[],
                metadata={},
                error="No files to merge"
            )
        
        # Parse first document
        merged = self.parse(file_paths[0])
        merged.metadata['merged_files'] = []
        
        # Merge additional documents
        for file_path in file_paths[1:]:
            doc = self.parse(file_path)
            if doc.success:
                merged.merge_with(doc)
        
        # Re-index tables
        for idx, table in enumerate(merged.tables):
            table.index = idx
        
        merged.metadata['total_files'] = len(file_paths)
        merged.metadata['merge_timestamp'] = datetime.now().isoformat()
        
        return merged
    
    def find_tables_with_pattern(self, document: ParsedDocument, pattern: str) -> List[Table]:
        """
        Find tables in document matching a header pattern
        
        Args:
            document: Parsed document
            pattern: Regex pattern to match headers
            
        Returns:
            List of matching tables
        """
        matching_tables = []
        regex = re.compile(pattern, re.IGNORECASE)
        
        for table in document.tables:
            headers_text = ' '.join(table.headers)
            if regex.search(headers_text):
                matching_tables.append(table)
        
        return matching_tables


# Convenience function for backward compatibility
def parse_document(file_path: Union[str, Path]) -> ParsedDocument:
    """Parse a document using the universal parser"""
    parser = UniversalDocumentParser()
    return parser.parse(file_path)