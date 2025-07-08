#!/usr/bin/env python3
"""
File Processing Service

Handles file reading, parsing, and preparation for knowledge graph construction.
Supports multiple file formats and batch processing.
"""

import os
import json
import asyncio
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import mimetypes

from core.logging import get_logger
from tools.base_service import BaseService

logger = get_logger(__name__)

class FileProcessor(BaseService):
    """
    Service for processing various file formats for knowledge graph construction.
    
    Supports: txt, md, json, csv, pdf, docx, and more
    """
    
    def __init__(self):
        """Initialize file processor."""
        super().__init__("FileProcessor")
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.xml': self._process_xml
        }
        
    async def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Process a single file for knowledge extraction.
        
        Args:
            file_path: Path to the file
            metadata: Optional metadata to include
            
        Returns:
            Processed file data ready for knowledge extraction
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_extension}")
                return await self._process_as_text(file_path)
            
            processor = self.supported_formats[file_extension]
            
            # Get file stats
            stat = file_path.stat()
            
            # Process file content
            content_data = await processor(file_path)
            
            # Build result
            result = {
                "file_id": str(file_path),
                "filename": file_path.name,
                "file_path": str(file_path.absolute()),
                "file_size": stat.st_size,
                "file_extension": file_extension,
                "mime_type": mimetypes.guess_type(str(file_path))[0],
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "processed_time": datetime.now().isoformat(),
                "metadata": metadata or {},
                **content_data
            }
            
            logger.info(f"Processed file: {file_path.name} ({len(result.get('content', ''))} chars)")
            return result
            
        except Exception as e:
            logger.error(f"File processing failed for {file_path}: {e}")
            return {
                "file_id": str(file_path),
                "filename": Path(file_path).name,
                "error": str(e),
                "content": "",
                "chunks": [],
                "processed_time": datetime.now().isoformat()
            }
    
    async def process_directory(self, 
                               directory_path: str, 
                               recursive: bool = True,
                               file_patterns: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Process all files in a directory.
        
        Args:
            directory_path: Directory to process
            recursive: Whether to process subdirectories
            file_patterns: File patterns to include (e.g., ['*.txt', '*.md'])
            
        Returns:
            List of processed file data
        """
        try:
            directory = Path(directory_path)
            if not directory.exists():
                raise FileNotFoundError(f"Directory not found: {directory_path}")
            
            # Find files
            files = []
            if recursive:
                for pattern in (file_patterns or ['*']):
                    files.extend(directory.rglob(pattern))
            else:
                for pattern in (file_patterns or ['*']):
                    files.extend(directory.glob(pattern))
            
            # Filter for files only and supported formats
            files = [f for f in files if f.is_file()]
            
            logger.info(f"Found {len(files)} files to process in {directory_path}")
            
            # Process files
            results = []
            for file_path in files:
                try:
                    result = await self.process_file(str(file_path))
                    results.append(result)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    results.append({
                        "file_id": str(file_path),
                        "filename": file_path.name,
                        "error": str(e)
                    })
            
            return results
            
        except Exception as e:
            logger.error(f"Directory processing failed: {e}")
            return []
    
    async def process_batch(self, file_paths: List[str], max_concurrent: int = 5) -> List[Dict[str, Any]]:
        """
        Process multiple files concurrently.
        
        Args:
            file_paths: List of file paths to process
            max_concurrent: Maximum concurrent file processing
            
        Returns:
            List of processed file data
        """
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_with_semaphore(file_path: str):
            async with semaphore:
                return await self.process_file(file_path)
        
        # Process files concurrently
        tasks = [process_with_semaphore(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Failed to process {file_paths[i]}: {result}")
                processed_results.append({
                    "file_id": file_paths[i],
                    "filename": Path(file_paths[i]).name,
                    "error": str(result)
                })
            else:
                processed_results.append(result)
        
        return processed_results
    
    # === Format-specific processors ===
    
    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            return {
                "content": content,
                "content_type": "text",
                "chunks": self._chunk_text(content),
                "word_count": len(content.split()),
                "line_count": len(content.splitlines())
            }
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Process Markdown files."""
        # For now, treat as text. Could add markdown parsing later
        return await self._process_text(file_path)
    
    async def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to text representation
            if isinstance(data, dict):
                content = self._json_to_text(data)
            elif isinstance(data, list):
                content = "\n".join(self._json_to_text(item) for item in data if isinstance(item, dict))
            else:
                content = str(data)
            
            return {
                "content": content,
                "content_type": "json",
                "chunks": self._chunk_text(content),
                "json_data": data,
                "json_keys": list(data.keys()) if isinstance(data, dict) else []
            }
        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files."""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n"
            content += f"Columns: {', '.join(df.columns)}\n\n"
            
            # Add sample rows as text
            for idx, row in df.head(100).iterrows():  # Limit to first 100 rows
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items())
                content += f"Row {idx + 1}: {row_text}\n"
            
            return {
                "content": content,
                "content_type": "csv",
                "chunks": self._chunk_text(content),
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns)
            }
        except ImportError:
            logger.warning("pandas not available for CSV processing, treating as text")
            return await self._process_text(file_path)
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files."""
        try:
            import PyPDF2
            
            content = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {e}")
            
            return {
                "content": content,
                "content_type": "pdf",
                "chunks": self._chunk_text(content),
                "page_count": len(reader.pages)
            }
        except ImportError:
            logger.warning("PyPDF2 not available for PDF processing")
            return {"content": "", "chunks": [], "error": "PyPDF2 not installed"}
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX files."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content += f"Table: {row_text}\n"
            
            return {
                "content": content,
                "content_type": "docx",
                "chunks": self._chunk_text(content),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except ImportError:
            logger.warning("python-docx not available for DOCX processing")
            return {"content": "", "chunks": [], "error": "python-docx not installed"}
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Process HTML files."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            return {
                "content": content,
                "content_type": "html",
                "chunks": self._chunk_text(content),
                "title": soup.title.string if soup.title else ""
            }
        except ImportError:
            logger.warning("beautifulsoup4 not available for HTML processing")
            return await self._process_text(file_path)
        except Exception as e:
            logger.error(f"HTML processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_xml(self, file_path: Path) -> Dict[str, Any]:
        """Process XML files."""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            content = self._xml_to_text(root)
            
            return {
                "content": content,
                "content_type": "xml",
                "chunks": self._chunk_text(content),
                "root_tag": root.tag
            }
        except Exception as e:
            logger.error(f"XML processing failed: {e}")
            return {"content": "", "chunks": [], "error": str(e)}
    
    async def _process_as_text(self, file_path: Path) -> Dict[str, Any]:
        """Fallback: process unknown format as text."""
        logger.info(f"Processing {file_path} as plain text")
        return await self._process_text(file_path)
    
    # === Helper methods ===
    
    def _chunk_text(self, text: str, chunk_size: int = 100000, overlap: int = 5000) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks. 
        
        Uses 100K chunk size to leverage modern LLM long context capabilities.
        Only splits when text exceeds 100K characters.
        """
        if len(text) <= chunk_size:
            return [{
                "chunk_id": 0,
                "content": text,
                "start": 0,
                "end": len(text),
                "size": len(text)
            }]
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            
            # Try to break at word boundary
            if end < len(text):
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk_content = text[start:end].strip()
            
            if chunk_content:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": chunk_content,
                    "start": start,
                    "end": end,
                    "size": len(chunk_content)
                })
                chunk_id += 1
            
            # Move start position with overlap
            if end - overlap > start:
                start = end - overlap
            else:
                start = end  # No overlap possible, move to next chunk
            
            # Prevent infinite loop  
            if start >= len(text):
                break
        
        return chunks
    
    def _json_to_text(self, data: Dict[str, Any], prefix: str = "") -> str:
        """Convert JSON object to readable text."""
        text_parts = []
        
        for key, value in data.items():
            if isinstance(value, dict):
                text_parts.append(f"{prefix}{key}:")
                text_parts.append(self._json_to_text(value, prefix + "  "))
            elif isinstance(value, list):
                text_parts.append(f"{prefix}{key}: [list with {len(value)} items]")
                for i, item in enumerate(value[:5]):  # Limit to first 5 items
                    if isinstance(item, dict):
                        text_parts.append(f"{prefix}  Item {i+1}:")
                        text_parts.append(self._json_to_text(item, prefix + "    "))
                    else:
                        text_parts.append(f"{prefix}  Item {i+1}: {item}")
            else:
                text_parts.append(f"{prefix}{key}: {value}")
        
        return "\n".join(text_parts)
    
    def _xml_to_text(self, element, level: int = 0) -> str:
        """Convert XML element to readable text."""
        indent = "  " * level
        text_parts = [f"{indent}{element.tag}"]
        
        if element.text and element.text.strip():
            text_parts.append(f"{indent}  {element.text.strip()}")
        
        for attr, value in element.attrib.items():
            text_parts.append(f"{indent}  @{attr}: {value}")
        
        for child in element:
            text_parts.append(self._xml_to_text(child, level + 1))
        
        return "\n".join(text_parts)

# Global instance
_file_processor = None

def get_file_processor() -> FileProcessor:
    """Get or create the global file processor instance."""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor()
    return _file_processor