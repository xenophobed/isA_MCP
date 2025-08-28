#!/usr/bin/env python3
"""
Document adapter for handling various document formats (PDF, DOC, DOCX, PPT, PPTX, TXT)
Provides document preprocessing, page extraction, and conversion to images.
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
from dataclasses import dataclass, field
from PIL import Image
import fitz  # PyMuPDF
import docx
# from pptx import Presentation  # TODO: Install python-pptx when needed
from pdf2image import convert_from_path
try:
    import pytesseract
    from PIL import Image as PILImage
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
import logging

from .base_adapter import FileAdapter
from .base_adapter import TableInfo, ColumnInfo 
from datetime import datetime

logger = logging.getLogger(__name__)

@dataclass
class DocumentPageInfo:
    """Information about a document page."""
    page_number: int
    page_type: str  # 'text', 'image', 'mixed', 'table'
    content: str
    image_path: Optional[str] = None
    tables: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class DocumentSection:
    """A logical section of a document."""
    section_id: str
    title: str
    content: str
    pages: List[int]
    section_type: str  # 'header', 'body', 'table', 'image', 'footer'
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ExtractedImage:
    """Information about an extracted image."""
    image_id: str
    page_number: int
    image_path: str
    bbox: Optional[Tuple[float, float, float, float]] = None
    image_type: str = 'embedded'  # 'embedded', 'figure', 'chart'
    caption: Optional[str] = None
    ocr_text: Optional[str] = None
    ocr_confidence: Optional[float] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class OCRResult:
    """OCR processing result."""
    image_id: str
    extracted_text: str
    confidence: float
    words: List[Dict[str, Any]] = field(default_factory=list)
    lines: List[Dict[str, Any]] = field(default_factory=list)
    paragraphs: List[Dict[str, Any]] = field(default_factory=list)

class DocumentAdapter(FileAdapter):
    """Adapter for processing various document formats."""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.doc': 'doc',
        '.docx': 'docx', 
        '.ppt': 'ppt',
        '.pptx': 'pptx',
        '.txt': 'txt'
    }
    
    def __init__(self):
        super().__init__()
        self.pages: List[DocumentPageInfo] = []
        self.sections: List[DocumentSection] = []
        self.extracted_images: List[ExtractedImage] = []
        self.temp_dir = None
        self.images_dir = None
        
    def _load_file(self, file_path: str) -> Any:
        """Load document file based on format."""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.SUPPORTED_FORMATS:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Create temporary directory for image outputs
            self.temp_dir = tempfile.mkdtemp(prefix='doc_analysis_')
            self.images_dir = os.path.join(self.temp_dir, 'images')
            os.makedirs(self.images_dir, exist_ok=True)
            
            format_type = self.SUPPORTED_FORMATS[file_ext]
            
            if format_type == 'pdf':
                return self._load_pdf(file_path)
            elif format_type == 'docx':
                return self._load_docx(file_path)
            elif format_type == 'pptx':
                return self._load_pptx(file_path)
            elif format_type == 'txt':
                return self._load_txt(file_path)
            else:
                raise ValueError(f"Format {format_type} not yet implemented")
                
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise
    
    def _load_pdf(self, file_path: str) -> fitz.Document:
        """Load PDF document."""
        doc = fitz.open(file_path)
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text
            text_content = page.get_text()
            
            # Extract tables
            tables = self._extract_tables_from_page(page)
            
            # Extract embedded images from page
            embedded_images = self._extract_images_from_page(page, page_num)
            
            # Convert page to image
            image_path = self._convert_page_to_image(page, page_num)
            
            # Determine page type
            page_type = self._determine_page_type(text_content, tables, page)
            
            page_info = DocumentPageInfo(
                page_number=page_num + 1,
                page_type=page_type,
                content=text_content,
                image_path=image_path,
                tables=tables,
                metadata={
                    'bbox': page.rect,
                    'rotation': page.rotation,
                    'image_count': len(page.get_images())
                }
            )
            
            self.pages.append(page_info)
        
        return doc
    
    def _load_docx(self, file_path: str) -> docx.Document:
        """Load DOCX document."""
        doc = docx.Document(file_path)
        
        # Process paragraphs and tables
        content_parts = []
        current_page = 1
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = docx.text.paragraph.Paragraph(element, doc)
                content_parts.append(para.text)
            elif element.tag.endswith('tbl'):  # Table
                table = docx.table.Table(element, doc)
                table_data = self._extract_docx_table(table)
                content_parts.append(f"[TABLE: {len(table_data)} rows]")
        
        # Create single page info for now (DOCX doesn't have clear page breaks)
        full_content = '\n'.join(content_parts)
        
        page_info = DocumentPageInfo(
            page_number=1,
            page_type='mixed',
            content=full_content,
            tables=[],
            metadata={'total_paragraphs': len(doc.paragraphs)}
        )
        
        self.pages.append(page_info)
        return doc
    
    def _load_pptx(self, file_path: str) -> Any:
        """Load PPTX presentation."""
        prs = Presentation(file_path)
        
        # Each slide becomes a page
        for slide_num, slide in enumerate(prs.slides):
            content_parts = []
            
            # Extract text from shapes
            for shape in slide.shapes:
                if hasattr(shape, 'text'):
                    content_parts.append(shape.text)
            
            # Convert slide to image
            image_path = self._convert_slide_to_image(slide, slide_num)
            
            page_info = DocumentPageInfo(
                page_number=slide_num + 1,
                page_type='mixed',
                content='\n'.join(content_parts),
                image_path=image_path,
                metadata={
                    'shape_count': len(slide.shapes),
                    'slide_layout': slide.slide_layout.name if hasattr(slide.slide_layout, 'name') else 'unknown'
                }
            )
            
            self.pages.append(page_info)
        
        return prs
    
    def _load_txt(self, file_path: str) -> str:
        """Load text file."""
        encoding = self._detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        
        # Split into logical pages (by double newlines or page breaks)
        pages = content.split('\n\n\n')  # Assuming triple newlines indicate page breaks
        
        for page_num, page_content in enumerate(pages):
            if page_content.strip():
                page_info = DocumentPageInfo(
                    page_number=page_num + 1,
                    page_type='text',
                    content=page_content.strip(),
                    metadata={'character_count': len(page_content)}
                )
                self.pages.append(page_info)
        
        return content
    
    def _extract_tables_from_page(self, page: fitz.Page) -> List[Dict[str, Any]]:
        """Extract tables from PDF page."""
        tables = []
        
        try:
            # Use PyMuPDF's table extraction
            tabs = page.find_tables()
            
            for i, tab in enumerate(tabs):
                table_data = tab.extract()
                if table_data:
                    tables.append({
                        'table_id': f'table_{i}',
                        'data': table_data,
                        'bbox': tab.bbox,
                        'rows': len(table_data),
                        'cols': len(table_data[0]) if table_data else 0
                    })
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
        
        return tables
    
    def _extract_images_from_page(self, page: fitz.Page, page_num: int) -> List[ExtractedImage]:
        """Extract embedded images from PDF page."""
        extracted_images = []
        
        try:
            # Get list of images on the page
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                # Get image reference
                xref = img[0]
                
                # Extract image data
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Save image to file
                image_filename = f"page_{page_num + 1}_img_{img_index + 1}.{image_ext}"
                image_path = os.path.join(self.images_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                
                # Get image bbox if available
                bbox = None
                try:
                    # Find image rectangle on page
                    for item in page.get_images(full=True):
                        if item[0] == xref:
                            # Get transform matrix and calculate bbox
                            bbox = item[1:]  # This is simplified - actual bbox extraction is more complex
                            break
                except Exception:
                    pass
                
                # Create ExtractedImage object
                extracted_img = ExtractedImage(
                    image_id=f"page_{page_num + 1}_img_{img_index + 1}",
                    page_number=page_num + 1,
                    image_path=image_path,
                    bbox=bbox,
                    image_type='embedded',
                    metadata={
                        'original_width': base_image.get("width"),
                        'original_height': base_image.get("height"),
                        'format': image_ext,
                        'size_bytes': len(image_bytes)
                    }
                )
                
                # Perform OCR if available and enabled
                if OCR_AVAILABLE:
                    ocr_result = self._perform_ocr_on_image(image_path)
                    if ocr_result:
                        extracted_img.ocr_text = ocr_result.extracted_text
                        extracted_img.ocr_confidence = ocr_result.confidence
                
                extracted_images.append(extracted_img)
                self.extracted_images.append(extracted_img)
                
        except Exception as e:
            logger.warning(f"Image extraction failed for page {page_num + 1}: {e}")
        
        return extracted_images
    
    def _perform_ocr_on_image(self, image_path: str) -> Optional[OCRResult]:
        """Perform OCR on an image file."""
        if not OCR_AVAILABLE:
            return None
        
        try:
            # Load image
            image = PILImage.open(image_path)
            
            # Perform OCR with detailed output
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, lang='eng')
            
            # Extract text
            extracted_text = pytesseract.image_to_string(image, lang='eng').strip()
            
            # Calculate average confidence
            confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            # Process word-level data
            words = []
            lines = []
            paragraphs = []
            
            for i, text in enumerate(ocr_data['text']):
                if text.strip():
                    word_info = {
                        'text': text,
                        'confidence': int(ocr_data['conf'][i]),
                        'bbox': (
                            ocr_data['left'][i],
                            ocr_data['top'][i], 
                            ocr_data['left'][i] + ocr_data['width'][i],
                            ocr_data['top'][i] + ocr_data['height'][i]
                        ),
                        'level': ocr_data['level'][i]
                    }
                    words.append(word_info)
            
            # Group words into lines and paragraphs
            current_line = []
            current_para = []
            last_line_num = None
            last_para_num = None
            
            for word in words:
                line_num = f"{ocr_data['page_num'][words.index(word)]}_{ocr_data['block_num'][words.index(word)]}_{ocr_data['par_num'][words.index(word)]}_{ocr_data['line_num'][words.index(word)]}"
                para_num = f"{ocr_data['page_num'][words.index(word)]}_{ocr_data['block_num'][words.index(word)]}_{ocr_data['par_num'][words.index(word)]}"
                
                if line_num != last_line_num:
                    if current_line:
                        lines.append({
                            'text': ' '.join([w['text'] for w in current_line]),
                            'confidence': sum([w['confidence'] for w in current_line]) / len(current_line),
                            'words': current_line
                        })
                    current_line = [word]
                    last_line_num = line_num
                else:
                    current_line.append(word)
                
                if para_num != last_para_num:
                    if current_para:
                        paragraphs.append({
                            'text': ' '.join([w['text'] for w in current_para]),
                            'confidence': sum([w['confidence'] for w in current_para]) / len(current_para),
                            'words': current_para
                        })
                    current_para = [word]
                    last_para_num = para_num
                else:
                    current_para.append(word)
            
            # Add last line and paragraph
            if current_line:
                lines.append({
                    'text': ' '.join([w['text'] for w in current_line]),
                    'confidence': sum([w['confidence'] for w in current_line]) / len(current_line),
                    'words': current_line
                })
            
            if current_para:
                paragraphs.append({
                    'text': ' '.join([w['text'] for w in current_para]),
                    'confidence': sum([w['confidence'] for w in current_para]) / len(current_para),
                    'words': current_para
                })
            
            return OCRResult(
                image_id=os.path.basename(image_path),
                extracted_text=extracted_text,
                confidence=avg_confidence / 100.0,  # Convert to 0-1 scale
                words=words,
                lines=lines,
                paragraphs=paragraphs
            )
            
        except Exception as e:
            logger.warning(f"OCR processing failed for {image_path}: {e}")
            return None
    
    def _convert_page_to_image(self, page: fitz.Page, page_num: int) -> str:
        """Convert PDF page to image."""
        try:
            # Render page as image
            mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)
            
            # Save as PNG
            image_path = os.path.join(self.images_dir, f'page_{page_num + 1}.png')
            pix.save(image_path)
            
            return image_path
        except Exception as e:
            logger.error(f"Failed to convert page {page_num} to image: {e}")
            return None
    
    def _convert_slide_to_image(self, slide, slide_num: int) -> str:
        """Convert PowerPoint slide to image."""
        # Note: This is a placeholder. Full implementation would require
        # converting PPTX to images using appropriate libraries
        image_path = os.path.join(self.images_dir, f'slide_{slide_num + 1}.png')
        
        # Create placeholder image for now
        img = Image.new('RGB', (800, 600), color='white')
        img.save(image_path)
        
        return image_path
    
    def _determine_page_type(self, text_content: str, tables: List[Dict], page: fitz.Page) -> str:
        """Determine the primary type of content on a page."""
        if not text_content.strip() and not tables:
            return 'image'
        elif tables and len(tables) >= 2:
            return 'table'
        elif len(text_content.strip()) < 100:
            return 'image'
        else:
            return 'mixed' if tables else 'text'
    
    def _extract_docx_table(self, table) -> List[List[str]]:
        """Extract data from DOCX table."""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            data.append(row_data)
        return data
    
    def _analyze_structure(self) -> Dict[str, Any]:
        """Analyze document structure."""
        return {
            'total_pages': len(self.pages),
            'page_types': {
                'text': len([p for p in self.pages if p.page_type == 'text']),
                'image': len([p for p in self.pages if p.page_type == 'image']),
                'table': len([p for p in self.pages if p.page_type == 'table']),
                'mixed': len([p for p in self.pages if p.page_type == 'mixed'])
            },
            'total_tables': sum(len(p.tables) for p in self.pages),
            'total_extracted_images': len(self.extracted_images),
            'has_images': any(p.image_path for p in self.pages),
            'temp_directory': self.temp_dir
        }
    
    def get_tables(self) -> List[TableInfo]:
        """Get table information from document."""
        tables = []
        
        for page in self.pages:
            for i, table_data in enumerate(page.tables):
                table_name = f"page_{page.page_number}_table_{i+1}"
                
                # Extract column information
                columns = []
                if table_data.get('data'):
                    headers = table_data['data'][0] if table_data['data'] else []
                    for j, header in enumerate(headers):
                        column_info = ColumnInfo(
                            table_name=table_name,
                            column_name=header if header.strip() else f"col_{j+1}",
                            data_type='TEXT',
                            is_nullable=True,
                            column_position=j+1
                        )
                        columns.append(column_info)
                
                table_info = TableInfo(
                    table_name=table_name,
                    schema_name='document',
                    table_type='EXTRACTED_TABLE',
                    record_count=table_data.get('rows', 0) - 1,  # Subtract header
                    table_comment=f'Extracted from page {page.page_number}',
                    created_date=datetime.now().isoformat(),
                    last_modified=datetime.now().isoformat()
                )
                # Store additional metadata as attributes
                table_info.extraction_metadata = {
                    'page_number': page.page_number,
                    'bbox': table_data.get('bbox'),
                    'extraction_method': 'document_adapter',
                    'columns': columns
                }
                tables.append(table_info)
        
        return tables
    
    def get_columns(self, table_name: str = None) -> List[ColumnInfo]:
        """Get column information."""
        columns = []
        tables = self.get_tables()
        
        for table in tables:
            if table_name is None or table.table_name == table_name:
                columns.extend(table.columns)
        
        return columns
    
    def get_page_images(self) -> List[str]:
        """Get list of page image paths."""
        return [p.image_path for p in self.pages if p.image_path]
    
    def get_extracted_images(self) -> List[ExtractedImage]:
        """Get list of extracted embedded images."""
        return self.extracted_images
    
    def get_images_by_page(self, page_number: int) -> List[ExtractedImage]:
        """Get extracted images for a specific page."""
        return [img for img in self.extracted_images if img.page_number == page_number]
    
    def get_image_by_id(self, image_id: str) -> Optional[ExtractedImage]:
        """Get specific image by ID."""
        for img in self.extracted_images:
            if img.image_id == image_id:
                return img
        return None
    
    def get_page_content(self, page_number: int) -> Optional[DocumentPageInfo]:
        """Get content for specific page."""
        for page in self.pages:
            if page.page_number == page_number:
                return page
        return None
    
    def classify_pages(self) -> Dict[str, List[int]]:
        """Classify pages by content type."""
        classification = {
            'text_heavy': [],
            'table_heavy': [],
            'image_heavy': [],
            'mixed': []
        }
        
        for page in self.pages:
            page_num = page.page_number
            
            if page.page_type == 'table' or len(page.tables) >= 2:
                classification['table_heavy'].append(page_num)
            elif page.page_type == 'text' and len(page.content) > 500:
                classification['text_heavy'].append(page_num)
            elif page.page_type == 'image' or len(page.content) < 50:
                classification['image_heavy'].append(page_num)
            else:
                classification['mixed'].append(page_num)
        
        return classification
    
    def _get_sample_data_internal(self, section_name: str, limit: int) -> List[Dict[str, Any]]:
        """Get sample data from document."""
        samples = []
        
        # Find matching page or table
        for page in self.pages[:limit]:
            if section_name.lower() in f"page_{page.page_number}".lower():
                samples.append({
                    'page': page.page_number,
                    'type': page.page_type,
                    'content_preview': page.content[:200] + '...' if len(page.content) > 200 else page.content,
                    'table_count': len(page.tables),
                    'has_image': page.image_path is not None
                })
        
        return samples
    
    def _analyze_column_distribution(self, section_name: str, column_name: str, sample_size: int) -> Dict[str, Any]:
        """Analyze column distribution in document tables."""
        return {
            'analysis_type': 'document_column',
            'section': section_name,
            'column': column_name,
            'note': 'Column distribution analysis for documents requires specific table extraction'
        }
    
    def _compute_quality_metrics(self, metrics: Dict[str, Any]) -> Dict[str, Any]:
        """Compute document quality metrics."""
        metrics.update({
            'total_sections': len(self.pages),
            'total_columns': sum(len(p.tables) for p in self.pages),
            'total_rows': sum(len(p.content.split('\n')) for p in self.pages),
            'pages_with_tables': len([p for p in self.pages if p.tables]),
            'pages_with_images': len([p for p in self.pages if p.image_path]),
            'average_content_length': sum(len(p.content) for p in self.pages) / len(self.pages) if self.pages else 0
        })
        
        return metrics
    
    def _get_data_summary(self) -> Dict[str, Any]:
        """Get summary of document data."""
        return {
            'total_pages': len(self.pages),
            'total_characters': sum(len(p.content) for p in self.pages),
            'total_tables': sum(len(p.tables) for p in self.pages),
            'page_types': [p.page_type for p in self.pages],
            'images_generated': len(self.get_page_images()),
            'temp_directory': self.temp_dir
        }
    
    def disconnect(self) -> None:
        """Clean up resources."""
        super().disconnect()
        
        # Clean up temporary files
        if self.temp_dir and os.path.exists(self.temp_dir):
            import shutil
            try:
                shutil.rmtree(self.temp_dir)
            except Exception as e:
                logger.warning(f"Failed to clean up temp directory {self.temp_dir}: {e}")
        
        self.pages.clear()
        self.sections.clear()
        self.extracted_images.clear()
        self.temp_dir = None
        self.images_dir = None